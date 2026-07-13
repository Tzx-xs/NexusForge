from __future__ import annotations

from dataclasses import dataclass
from enum import StrEnum
from html import escape as html_escape

from infrastructure.persistence.chapter_repo import ChapterRepository
from infrastructure.persistence.character_repo import CharacterRepository
from infrastructure.persistence.novel_repo import NovelRepository
from infrastructure.persistence.setting_repo import SettingRepository


class ExportFormat(StrEnum):
    TXT = "txt"
    MARKDOWN = "md"
    EPUB = "epub"
    DOCX = "docx"
    HTML = "html"


class ExportScope(StrEnum):
    FULL = "full"
    CHAPTER_RANGE = "chapter_range"
    SINGLE_CHAPTER = "single_chapter"
    OUTLINE_ONLY = "outline_only"


@dataclass
class ExportOptions:
    format: ExportFormat = ExportFormat.TXT
    scope: ExportScope = ExportScope.FULL
    include_title_page: bool = True
    include_chapter_numbers: bool = True
    include_toc: bool = True
    font_size: int = 14
    line_spacing: float = 1.5
    start_chapter: int | None = None
    end_chapter: int | None = None
    title: str | None = None
    author: str | None = None


class ExportService:
    """导出服务 - 支持多种格式的小说导出"""

    def __init__(
        self,
        novel_repo: NovelRepository,
        chapter_repo: ChapterRepository,
        character_repo: CharacterRepository | None = None,
        setting_repo: SettingRepository | None = None,
    ):
        self.novel_repo = novel_repo
        self.chapter_repo = chapter_repo
        self.character_repo = character_repo
        self.setting_repo = setting_repo

    def export_novel(self, novel_id: str, options: ExportOptions) -> bytes:
        novel = self.novel_repo.get_by_id(novel_id)
        if not novel:
            raise ValueError(f"Novel not found: {novel_id}")

        chapters = self._get_chapters(novel_id, options)

        if options.format == ExportFormat.TXT:
            return self._export_txt(novel, chapters, options)
        elif options.format == ExportFormat.MARKDOWN:
            return self._export_markdown(novel, chapters, options)
        elif options.format == ExportFormat.HTML:
            return self._export_html(novel, chapters, options)
        elif options.format == ExportFormat.DOCX:
            return self._export_docx(novel, chapters, options)
        elif options.format == ExportFormat.EPUB:
            return self._export_epub(novel, chapters, options)
        else:
            return self._export_txt(novel, chapters, options)

    def _get_chapters(self, novel_id: str, options: ExportOptions) -> list:
        all_chapters = self.chapter_repo.list_by_novel(novel_id)
        all_chapters.sort(key=lambda c: c.number or 0)

        if options.scope == ExportScope.SINGLE_CHAPTER and options.start_chapter:
            return [c for c in all_chapters if c.number == options.start_chapter]

        if options.scope == ExportScope.CHAPTER_RANGE:
            start = options.start_chapter or 1
            end = options.end_chapter or 9999
            return [c for c in all_chapters if start <= (c.number or 0) <= end]

        if options.scope == ExportScope.OUTLINE_ONLY:
            return []

        return all_chapters

    def _export_txt(self, novel, chapters, options: ExportOptions) -> bytes:
        lines = []

        title = options.title or novel.title
        author = options.author or "佚名"

        if options.include_title_page:
            lines.append(title)
            lines.append(f"作者：{author}")
            lines.append("")
            lines.append("=" * 40)
            lines.append("")

        if options.include_toc and len(chapters) > 0:
            lines.append("目录")
            lines.append("-" * 20)
            for ch in chapters:
                ch_num = f"第{ch.number}章 " if options.include_chapter_numbers and ch.number else ""
                lines.append(f"  {ch_num}{ch.title}")
            lines.append("")
            lines.append("=" * 40)
            lines.append("")

        for ch in chapters:
            ch_num = f"第{ch.number}章 " if options.include_chapter_numbers and ch.number else ""
            lines.append(f"{ch_num}{ch.title}")
            lines.append("")
            if ch.content:
                content = ch.content.strip()
                lines.append(content)
            lines.append("")
            lines.append("")

        return "\n".join(lines).encode("utf-8")

    def _export_markdown(self, novel, chapters, options: ExportOptions) -> bytes:
        lines = []

        title = options.title or novel.title
        author = options.author or "佚名"

        if options.include_title_page:
            lines.append(f"# {title}")
            lines.append("")
            lines.append(f"**作者：{author}**")
            lines.append("")
            lines.append("---")
            lines.append("")

        if novel.premise:
            lines.append(f"> {novel.premise}")
            lines.append("")

        if options.include_toc and len(chapters) > 0:
            lines.append("## 目录")
            lines.append("")
            for ch in chapters:
                ch_num = f"第{ch.number}章 " if options.include_chapter_numbers and ch.number else ""
                anchor = ch.title.lower().replace(" ", "-")
                lines.append(f"- [{ch_num}{ch.title}](#{anchor})")
            lines.append("")
            lines.append("---")
            lines.append("")

        for ch in chapters:
            ch_num = f"第{ch.number}章 " if options.include_chapter_numbers and ch.number else ""
            lines.append(f"## {ch_num}{ch.title}")
            lines.append("")
            if ch.content:
                paragraphs = [p.strip() for p in ch.content.split("\n") if p.strip()]
                for p in paragraphs:
                    lines.append(p)
                    lines.append("")
            lines.append("")

        return "\n".join(lines).encode("utf-8")

    def _export_html(self, novel, chapters, options: ExportOptions) -> bytes:
        # H-01 修复：所有用户输入（小说标题、作者、章节标题、章节正文）必须经 HTML 转义，
        # 防止存储型 XSS。安全实践：永远不直接拼接用户输入到 HTML 字符串。
        title = html_escape(options.title or novel.title)
        author = html_escape(options.author or "佚名")

        html_parts = [
            f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title}</title>
  <style>
    body {{
      font-family: "Songti SC", "SimSun", "Noto Serif CJK SC", serif;
      max-width: 720px;
      margin: 0 auto;
      padding: 40px 20px;
      line-height: {options.line_spacing};
      font-size: {options.font_size}px;
      color: #333;
    }}
    .title-page {{
      text-align: center;
      padding: 80px 0;
      border-bottom: 1px solid #eee;
      margin-bottom: 40px;
    }}
    .title-page h1 {{
      font-size: 2.5em;
      margin-bottom: 20px;
    }}
    .title-page .author {{
      font-size: 1.2em;
      color: #666;
    }}
    .toc {{
      margin-bottom: 40px;
    }}
    .toc h2 {{
      font-size: 1.5em;
      margin-bottom: 16px;
    }}
    .toc ul {{
      list-style: none;
      padding: 0;
    }}
    .toc li {{
      padding: 4px 0;
    }}
    .toc a {{
      color: #333;
      text-decoration: none;
    }}
    .toc a:hover {{
      color: #2080f0;
    }}
    .chapter {{
      margin-bottom: 48px;
    }}
    .chapter h2 {{
      font-size: 1.6em;
      margin-bottom: 20px;
      padding-bottom: 8px;
      border-bottom: 1px solid #eee;
    }}
    .chapter p {{
      text-indent: 2em;
      margin: 0 0 12px 0;
    }}
  </style>
</head>
<body>
"""
        ]

        if options.include_title_page:
            html_parts.append(f"""  <div class="title-page">
    <h1>{title}</h1>
    <div class="author">作者：{author}</div>
  </div>
""")

        if options.include_toc and len(chapters) > 0:
            html_parts.append('  <div class="toc">\n    <h2>目录</h2>\n    <ul>\n')
            for ch in chapters:
                ch_num = f"第{ch.number}章 " if options.include_chapter_numbers and ch.number else ""
                # 章节标题用户可编辑，必须转义
                safe_title = html_escape(ch.title)
                html_parts.append(f'      <li><a href="#ch-{ch.number or ch.id}">{ch_num}{safe_title}</a></li>\n')
            html_parts.append("    </ul>\n  </div>\n")

        for ch in chapters:
            ch_num = f"第{ch.number}章 " if options.include_chapter_numbers and ch.number else ""
            safe_title = html_escape(ch.title)
            html_parts.append(f'  <div class="chapter" id="ch-{ch.number or ch.id}">\n')
            html_parts.append(f"    <h2>{ch_num}{safe_title}</h2>\n")
            if ch.content:
                # 章节正文是核心用户输入，逐段转义
                paragraphs = [p.strip() for p in ch.content.split("\n") if p.strip()]
                for p in paragraphs:
                    html_parts.append(f"    <p>{html_escape(p)}</p>\n")
            html_parts.append("  </div>\n")

        html_parts.append("</body>\n</html>")
        return "".join(html_parts).encode("utf-8")

    def _export_docx(self, novel, chapters, options: ExportOptions) -> bytes:
        try:
            from io import BytesIO

            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Cm, Pt
        except ImportError:
            return self._export_txt(novel, chapters, options)

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(options.font_size)
        style.paragraph_format.line_spacing = options.line_spacing

        title = options.title or novel.title
        author = options.author or "佚名"

        if options.include_title_page:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(28)
            title_run.bold = True

            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"作者：{author}")
            author_run.font.size = Pt(16)

            doc.add_page_break()

        if options.include_toc and len(chapters) > 0:
            toc_heading = doc.add_heading("目录", level=1)
            toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for ch in chapters:
                ch_num = f"第{ch.number}章 " if options.include_chapter_numbers and ch.number else ""
                doc.add_paragraph(f"{ch_num}{ch.title}", style="List Bullet")

            doc.add_page_break()

        for ch in chapters:
            ch_num = f"第{ch.number}章 " if options.include_chapter_numbers and ch.number else ""
            doc.add_heading(f"{ch_num}{ch.title}", level=1)

            if ch.content:
                paragraphs = [p.strip() for p in ch.content.split("\n") if p.strip()]
                for p in paragraphs:
                    para = doc.add_paragraph()
                    para.paragraph_format.first_line_indent = Cm(0.74)
                    para.add_run(p)

            doc.add_paragraph()

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _export_epub(self, novel, chapters, options: ExportOptions) -> bytes:
        try:
            from io import BytesIO

            from ebooklib import epub
        except ImportError:
            return self._export_txt(novel, chapters, options)

        book = epub.EpubBook()

        title = options.title or novel.title
        author = options.author or "佚名"

        book.set_title(title)
        book.set_language("zh")
        book.add_author(author)

        book.spine = ["nav"]

        # EPUB 内 XHTML 拼接前对用户输入做 HTML/XML 转义，防止破坏标签结构
        safe_title = html_escape(title)
        safe_author = html_escape(author)

        if options.include_title_page:
            title_html = f"""<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{safe_title}</title>
    <style>
        body {{ text-align: center; padding-top: 100px; }}
        h1 {{ font-size: 28px; }}
        .author {{ font-size: 16px; color: #666; }}
    </style>
</head>
<body>
    <h1>{safe_title}</h1>
    <div class="author">作者：{safe_author}</div>
</body>
</html>"""
            title_page = epub.EpubHtml(title="封面", file_name="cover.xhtml")
            title_page.content = title_html
            book.add_item(title_page)
            book.spine.append(title_page)

        chapter_items = []
        for ch in chapters:
            ch_num = f"第{ch.number}章 " if options.include_chapter_numbers and ch.number else ""
            ch_title = f"{ch_num}{ch.title}"
            safe_ch_title = html_escape(ch_title)

            content_parts = [f"<h2>{safe_ch_title}</h2>"]
            if ch.content:
                paragraphs = [p.strip() for p in ch.content.split("\n") if p.strip()]
                for p in paragraphs:
                    content_parts.append(f"<p>{html_escape(p)}</p>")

            chapter_html = f"""<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{safe_ch_title}</title>
    <style>
        body {{ font-family: serif; line-height: 1.8; font-size: 16px; padding: 20px; }}
        h2 {{ font-size: 20px; border-bottom: 1px solid #eee; padding-bottom: 8px; }}
        p {{ text-indent: 2em; margin-bottom: 12px; }}
    </style>
</head>
<body>
{"".join(content_parts)}
</body>
</html>"""

            chapter_item = epub.EpubHtml(title=ch_title, file_name=f"chapter_{ch.number or ch.id}.xhtml")
            chapter_item.content = chapter_html
            book.add_item(chapter_item)
            chapter_items.append(chapter_item)
            book.spine.append(chapter_item)

        if options.include_toc and len(chapters) > 0:
            book.toc = [
                epub.Link(f"chapter_{ch.number or ch.id}.xhtml", f"{ch_num}{ch.title}", f"ch-{ch.id}")
                for ch_num, ch in [
                    (f"第{c.number}章 " if options.include_chapter_numbers and c.number else "", c) for c in chapters
                ]
            ]

        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        buffer = BytesIO()
        epub.write_epub(buffer, book, {})
        return buffer.getvalue()

    def get_export_filename(self, novel_id: str, options: ExportOptions) -> str:
        novel = self.novel_repo.get_by_id(novel_id)
        title = (options.title or (novel.title if novel else "novel")).replace(" ", "_")

        if options.scope == ExportScope.SINGLE_CHAPTER and options.start_chapter:
            scope_part = f"_ch{options.start_chapter}"
        elif options.scope == ExportScope.CHAPTER_RANGE:
            scope_part = f"_ch{options.start_chapter or 1}-{options.end_chapter or 'end'}"
        elif options.scope == ExportScope.OUTLINE_ONLY:
            scope_part = "_outline"
        else:
            scope_part = ""

        return f"{title}{scope_part}.{options.format.value}"
